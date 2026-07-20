"""
Document upload processing: text extraction, summarization, and chunk retrieval.
Supports multiple documents simultaneously with a hybrid strategy:
- Large tabular files (XLSX, CSV) → sent to Code Interpreter
- Text documents (PDF, TXT, DOCX) → summary + targeted chunk retrieval via session attributes
"""

import csv
import io
import logging
import re
import zipfile
from typing import Optional

import boto3
from pypdf import PdfReader
from docx import Document as DocxDocument
from openpyxl import load_workbook

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------
TOKEN_THRESHOLD = 80_000  # ~80K tokens ≈ 320K chars
CHAR_THRESHOLD = TOKEN_THRESHOLD * 4  # rough char estimate
CHUNK_SIZE = 2000  # characters per chunk for targeted retrieval
SESSION_ATTR_MAX_CHARS = 24_000  # ~25KB limit for promptSessionAttributes value
MAX_UPLOAD_FILES = 5  # Maximum number of files allowed at once

# Sensitivity labels that are NOT allowed (L3 and higher)
RESTRICTED_SENSITIVITY_LABELS = {"L3", "L4"}

# Classification keywords to search for in document headers/footers/text
# These are common classifications used in the Swiss federal administration
RESTRICTED_CLASSIFICATION_KEYWORDS = {
    "GEHEIM",
    "VERTRAULICH",
}


# ---------------------------------------------------------------------------
# Sensitivity Label Check (Microsoft Information Protection)
# ---------------------------------------------------------------------------

def check_sensitivity_label(file_bytes: bytes, filename: str) -> dict | None:
    """
    Check if a file contains a Microsoft Information Protection (MIP) sensitivity label.

    Supports:
    - PDF: reads raw bytes for MSIP_Label markers
    - DOCX/XLSX: reads custom XML parts inside the ZIP archive

    Returns a dict with label info if a label is found, e.g.:
        {"name": "L3", "guid": "...", "enabled": True}
    Returns None if no label is detected.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        return _check_sensitivity_label_pdf(file_bytes)
    elif ext in ("docx", "xlsx"):
        return _check_sensitivity_label_ooxml(file_bytes)
    else:
        # TXT and CSV don't carry MIP labels
        return None


def _check_sensitivity_label_pdf(file_bytes: bytes) -> dict | None:
    """Extract MSIP label from PDF metadata."""
    # Search for MSIP_Label_<GUID>_Name pattern in raw bytes
    match = re.search(
        rb"MSIP_Label_([0-9a-f\-]+)_Name[>\s/\(]+([A-Za-z0-9_\- ]+)",
        file_bytes,
    )
    if match:
        guid = match.group(1).decode("utf-8", errors="ignore")
        label_name = match.group(2).decode("utf-8", errors="ignore").strip().rstrip(")")
        return {"name": label_name, "guid": guid, "enabled": True}

    # Fallback: check XMP/pdfx namespace
    match = re.search(
        rb"MSIP_Label_([0-9a-f\-]+)_Name>([^<]+)<",
        file_bytes,
    )
    if match:
        guid = match.group(1).decode("utf-8", errors="ignore")
        label_name = match.group(2).decode("utf-8", errors="ignore").strip()
        return {"name": label_name, "guid": guid, "enabled": True}

    return None


def _check_sensitivity_label_ooxml(file_bytes: bytes) -> dict | None:
    """Extract MSIP label from Office Open XML (docx, xlsx) custom XML parts."""
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes), "r") as zf:
            for entry in zf.namelist():
                if "customXml/item" in entry or "docProps/custom" in entry:
                    try:
                        content = zf.read(entry).decode("utf-8", errors="ignore")
                    except Exception:
                        continue
                    if "MSIP_Label" not in content:
                        continue
                    # Look for the label name property
                    match = re.search(
                        r"MSIP_Label_([0-9a-f\-]+)_Name[^>]*>([^<]+)<",
                        content,
                    )
                    if match:
                        guid = match.group(1)
                        label_name = match.group(2).strip()
                        return {"name": label_name, "guid": guid, "enabled": True}
                    # Alternative: property with fmtid (custom.xml style)
                    match = re.search(
                        r'name="MSIP_Label_([0-9a-f\-]+)_Name"[^>]*>.*?<vt:lpwstr>([^<]+)',
                        content,
                        re.DOTALL,
                    )
                    if match:
                        guid = match.group(1)
                        label_name = match.group(2).strip()
                        return {"name": label_name, "guid": guid, "enabled": True}
    except zipfile.BadZipFile:
        logger.debug("File is not a valid ZIP/OOXML archive: %s", "check skipped")
    except Exception as e:
        logger.debug("Error checking sensitivity label in OOXML: %s", e)

    return None


def is_sensitivity_restricted(label_info: dict | None) -> bool:
    """Return True if the label indicates a restricted classification (L3+)."""
    if label_info is None:
        return False
    return label_info.get("name", "").upper() in {l.upper() for l in RESTRICTED_SENSITIVITY_LABELS}


def check_classification_in_text(file_bytes: bytes, filename: str) -> str | None:
    """
    Check if a document contains classification keywords (e.g. GEHEIM, VERTRAULICH, INTERN)
    in headers, footers, or the first/last lines of the document text.

    This catches documents from Acta Nova or other systems where no MIP label is set,
    but the classification is printed as text in the document.

    Returns the found keyword (e.g. "GEHEIM") or None.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        return _check_classification_text_pdf(file_bytes)
    elif ext == "docx":
        return _check_classification_text_docx(file_bytes)
    elif ext == "xlsx":
        # Excel files rarely have classification in headers as text
        return None
    else:
        return None


def _check_classification_text_pdf(file_bytes: bytes) -> str | None:
    """Check first and last page of a PDF for classification keywords."""
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        pages_to_check = []
        if reader.pages:
            pages_to_check.append(reader.pages[0])  # first page
        if len(reader.pages) > 1:
            pages_to_check.append(reader.pages[-1])  # last page

        for page in pages_to_check:
            text = (page.extract_text() or "").upper()
            # Check first ~500 and last ~500 chars (header/footer area)
            header_area = text[:500]
            footer_area = text[-500:] if len(text) > 500 else text
            for keyword in RESTRICTED_CLASSIFICATION_KEYWORDS:
                if keyword in header_area or keyword in footer_area:
                    return keyword
    except Exception as e:
        logger.debug("Error checking classification text in PDF: %s", e)
    return None


def _check_classification_text_docx(file_bytes: bytes) -> str | None:
    """Check headers, footers, and first paragraphs of a DOCX for classification keywords."""
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))

        # Check all section headers and footers
        for section in doc.sections:
            for header in (section.header, section.first_page_header):
                for para in header.paragraphs:
                    text_upper = para.text.strip().upper()
                    for keyword in RESTRICTED_CLASSIFICATION_KEYWORDS:
                        if keyword in text_upper:
                            return keyword
            for footer in (section.footer, section.first_page_footer):
                for para in footer.paragraphs:
                    text_upper = para.text.strip().upper()
                    for keyword in RESTRICTED_CLASSIFICATION_KEYWORDS:
                        if keyword in text_upper:
                            return keyword

        # Also check first few paragraphs (some docs have it as first line)
        for para in doc.paragraphs[:5]:
            text_upper = para.text.strip().upper()
            for keyword in RESTRICTED_CLASSIFICATION_KEYWORDS:
                if keyword in text_upper:
                    return keyword

    except Exception as e:
        logger.debug("Error checking classification text in DOCX: %s", e)
    return None


# ---------------------------------------------------------------------------
# Text Extraction
# ---------------------------------------------------------------------------

def extract_text_from_pdf(file_bytes: bytes) -> tuple[str, int]:
    """Extract text from a PDF file. Returns (text, page_count)."""
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        pages.append(text)
    full_text = "\n\n".join(pages)
    full_text = _clean_text(full_text)
    return full_text, len(reader.pages)


def extract_text_from_txt(file_bytes: bytes) -> tuple[str, int]:
    """Extract text from a plain text file. Returns (text, 1)."""
    text = file_bytes.decode("utf-8", errors="replace")
    text = _clean_text(text)
    return text, 1


def extract_text_from_docx(file_bytes: bytes) -> tuple[str, int]:
    """Extract text from a DOCX file. Returns (text, page_count_estimate)."""
    doc = DocxDocument(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n\n".join(paragraphs)
    full_text = _clean_text(full_text)
    # Estimate page count (~3000 chars per page)
    page_estimate = max(1, len(full_text) // 3000)
    return full_text, page_estimate


def extract_text_from_xlsx(file_bytes: bytes) -> tuple[str, int]:
    """Extract text from an Excel (.xlsx) file. Returns (text, sheet_count)."""
    wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    sheets_text = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cell_values = [str(cell) if cell is not None else "" for cell in row]
            # Skip completely empty rows
            if any(v.strip() for v in cell_values):
                rows.append(" | ".join(cell_values))
        if rows:
            header = f"--- Sheet: {sheet_name} ---"
            sheets_text.append(f"{header}\n" + "\n".join(rows))
    wb.close()
    full_text = "\n\n".join(sheets_text)
    full_text = _clean_text(full_text)
    return full_text, len(wb.sheetnames)


def extract_text_from_csv(file_bytes: bytes) -> tuple[str, int]:
    """Extract text from a CSV file. Returns (text, 1)."""
    text_content = file_bytes.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text_content))
    rows = []
    for row in reader:
        if any(cell.strip() for cell in row):
            rows.append(" | ".join(row))
    full_text = "\n".join(rows)
    full_text = _clean_text(full_text)
    return full_text, 1


def extract_text(file_bytes: bytes, filename: str) -> tuple[str, int]:
    """
    Extract text from an uploaded file based on its extension.
    Returns (extracted_text, page_count).
    Raises ValueError if the file type is unsupported or extraction fails.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        text, pages = extract_text_from_pdf(file_bytes)
    elif ext == "txt":
        text, pages = extract_text_from_txt(file_bytes)
    elif ext == "docx":
        text, pages = extract_text_from_docx(file_bytes)
    elif ext == "xlsx":
        text, pages = extract_text_from_xlsx(file_bytes)
    elif ext == "csv":
        text, pages = extract_text_from_csv(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: .{ext}")

    # Skip quality check for tabular files (often mostly numeric)
    if ext not in ("xlsx", "csv") and not text_quality_ok(text):
        raise ValueError(
            "The extracted text appears to be of low quality (possibly a scanned PDF). "
            "Please upload a text-based document."
        )

    return text, pages


# ---------------------------------------------------------------------------
# Text Quality & Cleaning
# ---------------------------------------------------------------------------

def _clean_text(text: str) -> str:
    """Remove SVG artifacts and normalize whitespace."""
    text = remove_svg_artifacts(text)
    # Collapse multiple blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def remove_svg_artifacts(text: str) -> str:
    """Remove SVG/XML-like artifacts from extracted text."""
    text = re.sub(r"<svg[^>]*>.*?</svg>", "", text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<[^>]+>", "", text)
    return text


def text_quality_ok(text: str, min_alpha_ratio: float = 0.3, min_length: int = 50) -> bool:
    """Check if extracted text has acceptable quality."""
    if len(text.strip()) < min_length:
        return False
    alpha_chars = sum(1 for c in text if c.isalpha())
    total_chars = len(text.replace(" ", "").replace("\n", ""))
    if total_chars == 0:
        return False
    return (alpha_chars / total_chars) >= min_alpha_ratio


# ---------------------------------------------------------------------------
# Summarization
# ---------------------------------------------------------------------------

def summarize_document(text: str, region: Optional[str] = None) -> str:
    """
    Summarize a large document using Bedrock Claude.
    Returns a comprehensive summary preserving key facts and structure.
    """
    import os
    import json

    region = region or os.getenv("AWS_REGION", "us-east-1")
    bedrock_runtime = boto3.client("bedrock-runtime", region_name=region)

    # Truncate input to ~200K chars to stay within model limits
    max_input_chars = 200_000
    truncated = text[:max_input_chars]
    if len(text) > max_input_chars:
        truncated += "\n\n[... document truncated for summarization ...]"

    prompt = (
        "Summarize the following document comprehensively, preserving key facts, "
        "figures, dates, names, and structural sections. The summary should be "
        "detailed enough to answer most questions about the document's content.\n\n"
        f"DOCUMENT:\n{truncated}"
    )

    body = json.dumps({
        "anthropic_version": "bedrock-2023-05-31",
        "max_tokens": 4096,
        "messages": [{"role": "user", "content": prompt}],
    })

    try:
        response = bedrock_runtime.invoke_model(
            modelId="anthropic.claude-3-haiku-20240307-v1:0",
            contentType="application/json",
            accept="application/json",
            body=body,
        )
        result = json.loads(response["body"].read())
        return result["content"][0]["text"]
    except Exception as e:
        logger.error("Summarization failed: %s", e)
        # Fallback: return truncated beginning
        fallback_len = SESSION_ATTR_MAX_CHARS - 500
        return text[:fallback_len] + "\n\n[... truncated, summarization unavailable ...]"


# ---------------------------------------------------------------------------
# Smart Context Strategy
# ---------------------------------------------------------------------------

def prepare_document_context(extracted_text: str, file_ext: str = "") -> tuple[str, str]:
    """
    Decide whether to use full text, a summary, or Code Interpreter based on size and file type.
    Returns (context_text, context_mode) where context_mode is "full", "summary", or "code_interpreter".

    For tabular files (xlsx, csv), summarization is skipped — large tables are
    routed to Code Interpreter by the caller.
    """
    is_tabular = file_ext in ("xlsx", "csv")

    if is_tabular:
        # For tables: if small enough, send full text; otherwise flag as oversized
        if len(extracted_text) <= SESSION_ATTR_MAX_CHARS:
            return extracted_text, "full"
        else:
            # Caller will route this to Code Interpreter
            note = (
                f"[Large tabular document — {len(extracted_text):,} characters. "
                f"File will be sent to Code Interpreter for analysis.]"
            )
            return note, "code_interpreter"

    # Non-tabular documents: existing logic
    if len(extracted_text) < CHAR_THRESHOLD:
        if len(extracted_text) <= SESSION_ATTR_MAX_CHARS:
            return extracted_text, "full"
        else:
            summary = summarize_document(extracted_text)
            return summary, "summary"
    else:
        summary = summarize_document(extracted_text)
        return summary, "summary"


# ---------------------------------------------------------------------------
# Targeted Chunk Retrieval (for large documents)
# ---------------------------------------------------------------------------

def chunk_text(text: str, chunk_size: int = CHUNK_SIZE) -> list[str]:
    """Split text into overlapping chunks."""
    chunks = []
    step = chunk_size - 200  # 200-char overlap
    for i in range(0, len(text), step):
        chunk = text[i:i + chunk_size]
        if chunk.strip():
            chunks.append(chunk)
    return chunks


def find_relevant_chunks(full_text: str, query: str, top_k: int = 3, is_tabular: bool = False) -> str:
    """
    Simple keyword-based chunk retrieval.
    Returns the top matching chunks concatenated.
    """
    chunks = chunk_text(full_text)

    if not chunks:
        return ""

    # Extract keywords from query (words > 1 char)
    stopwords = {
        # German
        "der", "die", "das", "und", "oder", "ein", "eine", "ist", "sind", "von", "für", "mit", "auf",
        "den", "dem", "des", "sich", "auch", "wird", "hat", "kann", "als", "bei", "noch", "nach",
        # French
        "les", "des", "une", "est", "sont", "par", "pour", "dans", "sur", "avec", "qui", "que",
        "pas", "plus", "aux", "son", "ses", "cette", "ces", "ont", "été",
        # Italian
        "gli", "una", "del", "dei", "che", "per", "con", "sono", "nel", "nella", "della",
        "suo", "suoi", "questa", "questi", "alle", "anche", "più", "non",
        # English
        "the", "and", "for", "with", "from", "that", "this", "are", "was", "were", "been", "have",
        "has", "not", "but", "can", "will",
    }
    keywords = [w.lower() for w in re.split(r"\W+", query) if len(w) > 1 and w.lower() not in stopwords]
    if not keywords:
        return ""

    # Score each chunk by keyword occurrence
    scored = []
    for chunk in chunks:
        chunk_lower = chunk.lower()
        score = sum(chunk_lower.count(kw) for kw in keywords)
        if score > 0:
            scored.append((score, chunk))

    scored.sort(key=lambda x: x[0], reverse=True)
    top_chunks = [chunk for _, chunk in scored[:top_k]]

    if not top_chunks:
        return ""

    result = "\n\n---\n\n".join(top_chunks)
    # Ensure it fits in session attributes
    if len(result) > SESSION_ATTR_MAX_CHARS:
        result = result[:SESSION_ATTR_MAX_CHARS]
    return result


# ---------------------------------------------------------------------------
# Multi-Document Processing
# ---------------------------------------------------------------------------

def process_multiple_documents(files_data: list[dict]) -> dict:
    """
    Process multiple uploaded documents and categorize them by handling strategy.

    Args:
        files_data: List of dicts with keys: 'name', 'bytes'

    Returns:
        Dict with:
        - 'text_docs': list of dicts with 'name', 'full_text', 'page_count', 'context', 'context_mode'
        - 'code_interpreter_docs': list of dicts with 'name', 'bytes', 'media_type'
        - 'errors': list of dicts with 'name', 'error'
    """
    result = {
        "text_docs": [],
        "code_interpreter_docs": [],
        "errors": [],
    }

    for file_info in files_data:
        name = file_info["name"]
        file_bytes = file_info["bytes"]
        ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""

        # Check sensitivity label BEFORE processing
        label_info = check_sensitivity_label(file_bytes, name)
        if is_sensitivity_restricted(label_info):
            label_name = label_info.get("name", "unbekannt")
            result["errors"].append({
                "name": name,
                "error": (
                    f"⚠️ Dieses Dokument ist als «{label_name}» klassifiziert. "
                    f"Vertrauliche und geheime Dokumente (L3 und höher) dürfen nicht "
                    f"im Chatbot verarbeitet werden. Bitte laden Sie nur öffentliche "
                    f"oder interne Dokumente hoch."
                ),
                "sensitivity_blocked": True,
            })
            continue

        # Fallback: check for classification keywords in document text (headers/footers)
        classification_keyword = check_classification_in_text(file_bytes, name)
        if classification_keyword:
            result["errors"].append({
                "name": name,
                "error": (
                    f"⚠️ Dieses Dokument enthält die Klassifizierung «{classification_keyword}». "
                    f"Vertrauliche und geheime Dokumente dürfen nicht "
                    f"im Chatbot verarbeitet werden. Bitte laden Sie nur öffentliche "
                    f"oder interne Dokumente hoch."
                ),
                "sensitivity_blocked": True,
            })
            continue

        try:
            extracted_text, page_count = extract_text(file_bytes, name)
        except (ValueError, Exception) as e:
            result["errors"].append({"name": name, "error": str(e)})
            continue

        # Decide routing based on file type and size
        is_tabular = ext in ("xlsx", "csv")

        if is_tabular and len(extracted_text) > SESSION_ATTR_MAX_CHARS:
            # Large table → Code Interpreter
            if ext == "csv":
                media_type = "text/csv"
            else:
                media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            result["code_interpreter_docs"].append({
                "name": name,
                "bytes": file_bytes,
                "media_type": media_type,
            })
        else:
            # Text-based or small tabular → session attributes path
            doc_context, context_mode = prepare_document_context(extracted_text, file_ext=ext)
            result["text_docs"].append({
                "name": name,
                "full_text": extracted_text,
                "page_count": page_count,
                "context": doc_context,
                "context_mode": context_mode,
            })

    return result


def build_multi_doc_context(text_docs: list[dict], query: str = "") -> str:
    """
    Build a combined context string from multiple text documents for session attributes.
    Uses summaries + targeted chunk retrieval to stay within size limits.

    Args:
        text_docs: List from process_multiple_documents()['text_docs']
        query: The user's current question (for targeted retrieval)

    Returns:
        Combined context string within SESSION_ATTR_MAX_CHARS
    """
    if not text_docs:
        return ""

    # Single document — use existing logic directly
    if len(text_docs) == 1:
        doc = text_docs[0]
        context = doc["context"]
        if doc["context_mode"] == "summary" and query and doc.get("full_text"):
            relevant_chunks = find_relevant_chunks(doc["full_text"], query, is_tabular=False)
            if relevant_chunks:
                context = (
                    f"DOCUMENT SUMMARY:\n{context}\n\n"
                    f"RELEVANT SECTIONS:\n{relevant_chunks}"
                )
        return context

    # Multiple documents — allocate space proportionally
    budget_per_doc = SESSION_ATTR_MAX_CHARS // len(text_docs)
    # Reserve some space for headers
    budget_per_doc = max(budget_per_doc - 200, 1000)

    parts = []
    for doc in text_docs:
        header = f"=== DOCUMENT: {doc['name']} ===\n"
        context = doc["context"]

        # For summarized docs, try to add relevant chunks
        if doc["context_mode"] == "summary" and query and doc.get("full_text"):
            relevant_chunks = find_relevant_chunks(
                doc["full_text"], query, top_k=2, is_tabular=False
            )
            if relevant_chunks:
                context = (
                    f"SUMMARY:\n{context}\n\n"
                    f"RELEVANT SECTIONS:\n{relevant_chunks}"
                )

        # Truncate to budget
        if len(context) > budget_per_doc:
            context = context[:budget_per_doc] + "\n[... truncated ...]"

        parts.append(header + context)

    combined = "\n\n".join(parts)

    # Final safety truncation
    if len(combined) > SESSION_ATTR_MAX_CHARS:
        combined = combined[:SESSION_ATTR_MAX_CHARS]

    return combined


def build_code_interpreter_files(code_interpreter_docs: list[dict]) -> list[dict]:
    """
    Build the files list for Bedrock Agent's sessionState.files from Code Interpreter documents.

    Args:
        code_interpreter_docs: List from process_multiple_documents()['code_interpreter_docs']

    Returns:
        List of file dicts ready for sessionState['files']
    """
    agent_files = []
    for doc in code_interpreter_docs:
        agent_files.append({
            "name": doc["name"],
            "source": {
                "sourceType": "BYTE_CONTENT",
                "byteContent": {
                    "data": doc["bytes"],
                    "mediaType": doc["media_type"],
                }
            },
            "useCase": "CODE_INTERPRETER",
        })
    return agent_files
