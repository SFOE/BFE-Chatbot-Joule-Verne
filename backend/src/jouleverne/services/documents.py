"""Document processing — text extraction, summarization, and chunking.

Ported from the Streamlit app's src/document_processing.py.
Supports multiple documents with a hybrid strategy:
- Large tabular files (XLSX, CSV) → flagged for Code Interpreter
- Text documents (PDF, TXT, DOCX) → full text or summary via session attributes
"""

import csv
import io
import json
import logging
import re
import zipfile
from typing import Optional

import boto3
from pypdf import PdfReader
from docx import Document as DocxDocument
from openpyxl import load_workbook

from ..config import settings

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------
TOKEN_THRESHOLD = 80_000  # ~80K tokens ≈ 320K chars
CHAR_THRESHOLD = TOKEN_THRESHOLD * 4
CHUNK_SIZE = 2000
SESSION_ATTR_MAX_CHARS = 24_000
MAX_UPLOAD_FILES = 5
MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB

ALLOWED_EXTENSIONS = {"pdf", "txt", "docx", "xlsx", "csv"}
MEDIA_TYPES = {
    "csv": "text/csv",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}

# ---------------------------------------------------------------------------
# Sensitivity / Classification
# ---------------------------------------------------------------------------

# Sensitivity labels that are NOT allowed (L3 and higher)
RESTRICTED_SENSITIVITY_LABELS = {"L3", "L4"}

# Classification keywords to search for in document headers/footers/text
# These are common classifications used in the Swiss federal administration
RESTRICTED_CLASSIFICATION_KEYWORDS = {
    "GEHEIM",
    "VERTRAULICH",
}


def check_sensitivity_label(file_bytes: bytes, filename: str) -> dict | None:
    """Check if a file contains a Microsoft Information Protection (MIP) sensitivity label.

    Supports:
    - PDF: reads raw bytes for MSIP_Label markers
    - DOCX/XLSX: reads custom XML parts inside the ZIP archive

    Returns a dict with label info if found, e.g.:
        {"name": "L3", "guid": "...", "enabled": True}
    Returns None if no label is detected.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        return _check_sensitivity_label_pdf(file_bytes)
    elif ext in ("docx", "xlsx"):
        return _check_sensitivity_label_ooxml(file_bytes)
    else:
        return None


def _check_sensitivity_label_pdf(file_bytes: bytes) -> dict | None:
    """Extract MSIP label from PDF metadata."""
    match = re.search(
        rb"MSIP_Label_([0-9a-f\-]+)_Name[>\s/\(]+([A-Za-z0-9_\- ]+)",
        file_bytes,
    )
    if match:
        guid = match.group(1).decode("utf-8", errors="ignore")
        label_name = match.group(2).decode("utf-8", errors="ignore").strip().rstrip(")")
        return {"name": label_name, "guid": guid, "enabled": True}

    # Fallback: check XMP/pdfx namespace
    match = re.search(
        rb"MSIP_Label_([0-9a-f\-]+)_Name>([^<]+)<",
        file_bytes,
    )
    if match:
        guid = match.group(1).decode("utf-8", errors="ignore")
        label_name = match.group(2).decode("utf-8", errors="ignore").strip()
        return {"name": label_name, "guid": guid, "enabled": True}

    return None


def _check_sensitivity_label_ooxml(file_bytes: bytes) -> dict | None:
    """Extract MSIP label from Office Open XML (docx, xlsx) custom XML parts."""
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes), "r") as zf:
            for entry in zf.namelist():
                if "customXml/item" in entry or "docProps/custom" in entry:
                    try:
                        content = zf.read(entry).decode("utf-8", errors="ignore")
                    except Exception:
                        continue
                    if "MSIP_Label" not in content:
                        continue
                    match = re.search(
                        r"MSIP_Label_([0-9a-f\-]+)_Name[^>]*>([^<]+)<",
                        content,
                    )
                    if match:
                        guid = match.group(1)
                        label_name = match.group(2).strip()
                        return {"name": label_name, "guid": guid, "enabled": True}
                    # Alternative: property with fmtid (custom.xml style)
                    match = re.search(
                        r'name="MSIP_Label_([0-9a-f\-]+)_Name"[^>]*>.*?<vt:lpwstr>([^<]+)',
                        content,
                        re.DOTALL,
                    )
                    if match:
                        guid = match.group(1)
                        label_name = match.group(2).strip()
                        return {"name": label_name, "guid": guid, "enabled": True}
    except zipfile.BadZipFile:
        logger.debug("File is not a valid ZIP/OOXML archive: check skipped")
    except Exception as e:
        logger.debug("Error checking sensitivity label in OOXML: %s", e)

    return None


def is_sensitivity_restricted(label_info: dict | None) -> bool:
    """Return True if the label indicates a restricted classification (L3+)."""
    if label_info is None:
        return False
    return label_info.get("name", "").upper() in {label.upper() for label in RESTRICTED_SENSITIVITY_LABELS}


def check_classification_in_text(file_bytes: bytes, filename: str) -> str | None:
    """Check if a document contains classification keywords (e.g. GEHEIM, VERTRAULICH)
    in headers, footers, or the first/last lines of the document text.

    Returns the found keyword or None.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        return _check_classification_text_pdf(file_bytes)
    elif ext == "docx":
        return _check_classification_text_docx(file_bytes)
    elif ext == "xlsx":
        return None
    else:
        return None


def _check_classification_text_pdf(file_bytes: bytes) -> str | None:
    """Check first and last page of a PDF for classification keywords."""
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        pages_to_check = []
        if reader.pages:
            pages_to_check.append(reader.pages[0])
        if len(reader.pages) > 1:
            pages_to_check.append(reader.pages[-1])

        for page in pages_to_check:
            text = (page.extract_text() or "").upper()
            header_area = text[:500]
            footer_area = text[-500:] if len(text) > 500 else text
            for keyword in RESTRICTED_CLASSIFICATION_KEYWORDS:
                if keyword in header_area or keyword in footer_area:
                    return keyword
    except Exception as e:
        logger.debug("Error checking classification text in PDF: %s", e)
    return None


def _check_classification_text_docx(file_bytes: bytes) -> str | None:
    """Check headers, footers, and first paragraphs of a DOCX for classification keywords."""
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))

        for section in doc.sections:
            for header in (section.header, section.first_page_header):
                for para in header.paragraphs:
                    text_upper = para.text.strip().upper()
                    for keyword in RESTRICTED_CLASSIFICATION_KEYWORDS:
                        if keyword in text_upper:
                            return keyword
            for footer in (section.footer, section.first_page_footer):
                for para in footer.paragraphs:
                    text_upper = para.text.strip().upper()
                    for keyword in RESTRICTED_CLASSIFICATION_KEYWORDS:
                        if keyword in text_upper:
                            return keyword

        for para in doc.paragraphs[:5]:
            text_upper = para.text.strip().upper()
            for keyword in RESTRICTED_CLASSIFICATION_KEYWORDS:
                if keyword in text_upper:
                    return keyword

    except Exception as e:
        logger.debug("Error checking classification text in DOCX: %s", e)
    return None


# ---------------------------------------------------------------------------
# Text Extraction
# ---------------------------------------------------------------------------

def extract_text(file_bytes: bytes, filename: str) -> tuple[str, int]:
    """Extract text from an uploaded file. Returns (text, page_count).

    Raises ValueError for unsupported types or low-quality extraction.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    extractors = {
        "pdf": _extract_pdf,
        "txt": _extract_txt,
        "docx": _extract_docx,
        "xlsx": _extract_xlsx,
        "csv": _extract_csv,
    }

    extractor = extractors.get(ext)
    if not extractor:
        raise ValueError(f"Unsupported file type: .{ext}")

    text, pages = extractor(file_bytes)

    # Skip quality check for tabular files
    if ext not in ("xlsx", "csv") and not _text_quality_ok(text):
        raise ValueError(
            "Extracted text appears low quality (possibly a scanned PDF). "
            "Please upload a text-based document."
        )

    return text, pages


def _extract_pdf(file_bytes: bytes) -> tuple[str, int]:
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = [page.extract_text() or "" for page in reader.pages]
    return _clean_text("\n\n".join(pages)), len(reader.pages)


def _extract_txt(file_bytes: bytes) -> tuple[str, int]:
    text = file_bytes.decode("utf-8", errors="replace")
    return _clean_text(text), 1


def _extract_docx(file_bytes: bytes) -> tuple[str, int]:
    doc = DocxDocument(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = _clean_text("\n\n".join(paragraphs))
    page_estimate = max(1, len(full_text) // 3000)
    return full_text, page_estimate


def _extract_xlsx(file_bytes: bytes) -> tuple[str, int]:
    wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    sheets_text = []
    sheet_count = len(wb.sheetnames)
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cell_values = [str(cell) if cell is not None else "" for cell in row]
            if any(v.strip() for v in cell_values):
                rows.append(" | ".join(cell_values))
        if rows:
            sheets_text.append(f"--- Sheet: {sheet_name} ---\n" + "\n".join(rows))
    wb.close()
    return _clean_text("\n\n".join(sheets_text)), sheet_count


def _extract_csv(file_bytes: bytes) -> tuple[str, int]:
    text_content = file_bytes.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text_content))
    rows = [" | ".join(row) for row in reader if any(cell.strip() for cell in row)]
    return _clean_text("\n".join(rows)), 1


# ---------------------------------------------------------------------------
# Text Quality & Cleaning
# ---------------------------------------------------------------------------

def _clean_text(text: str) -> str:
    text = re.sub(r"<svg[^>]*>.*?</svg>", "", text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<[^>]+>", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _text_quality_ok(text: str, min_alpha_ratio: float = 0.3, min_length: int = 50) -> bool:
    if len(text.strip()) < min_length:
        return False
    alpha_chars = sum(1 for c in text if c.isalpha())
    total_chars = len(text.replace(" ", "").replace("\n", ""))
    if total_chars == 0:
        return False
    return (alpha_chars / total_chars) >= min_alpha_ratio


# ---------------------------------------------------------------------------
# Summarization
# ---------------------------------------------------------------------------

def summarize_document(text: str) -> str:
    """Summarize a large document using Bedrock Claude Haiku."""
    bedrock_runtime = boto3.client("bedrock-runtime", region_name=settings.AWS_REGION)

    max_input_chars = 200_000
    truncated = text[:max_input_chars]
    if len(text) > max_input_chars:
        truncated += "\n\n[... document truncated for summarization ...]"

    prompt = (
        "Summarize the following document comprehensively, preserving key facts, "
        "figures, dates, names, and structural sections. The summary should be "
        "detailed enough to answer most questions about the document's content.\n\n"
        f"DOCUMENT:\n{truncated}"
    )

    body = json.dumps({
        "anthropic_version": "bedrock-2023-05-31",
        "max_tokens": 4096,
        "messages": [{"role": "user", "content": prompt}],
    })

    try:
        response = bedrock_runtime.invoke_model(
            modelId="anthropic.claude-3-haiku-20240307-v1:0",
            contentType="application/json",
            accept="application/json",
            body=body,
        )
        result = json.loads(response["body"].read())
        return result["content"][0]["text"]
    except Exception as e:
        logger.error("Summarization failed: %s", e)
        fallback_len = SESSION_ATTR_MAX_CHARS - 500
        return text[:fallback_len] + "\n\n[... truncated, summarization unavailable ...]"


# ---------------------------------------------------------------------------
# Context Preparation
# ---------------------------------------------------------------------------

def prepare_document_context(extracted_text: str, file_ext: str = "") -> tuple[str, str]:
    """Decide whether to use full text or summary.

    Returns (context_text, context_mode) where context_mode is
    "full", "summary", or "code_interpreter".
    """
    is_tabular = file_ext in ("xlsx", "csv")

    if is_tabular:
        if len(extracted_text) <= SESSION_ATTR_MAX_CHARS:
            return extracted_text, "full"
        else:
            return (
                f"[Large tabular document — {len(extracted_text):,} characters. "
                f"File will be sent to Code Interpreter for analysis.]"
            ), "code_interpreter"

    if len(extracted_text) <= SESSION_ATTR_MAX_CHARS:
        return extracted_text, "full"
    else:
        summary = summarize_document(extracted_text)
        return summary, "summary"


# ---------------------------------------------------------------------------
# Multi-Document Processing
# ---------------------------------------------------------------------------

def process_multiple_documents(files_data: list[dict]) -> dict:
    """Process uploaded documents and categorize by handling strategy.

    Args:
        files_data: List of dicts with 'name' and 'bytes' keys.

    Returns:
        Dict with 'text_docs', 'code_interpreter_docs', and 'errors' lists.
    """
    result = {
        "text_docs": [],
        "code_interpreter_docs": [],
        "errors": [],
    }

    for file_info in files_data:
        name = file_info["name"]
        file_bytes = file_info["bytes"]
        ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""

        # --- Sensitivity check BEFORE processing ---
        label_info = check_sensitivity_label(file_bytes, name)
        if is_sensitivity_restricted(label_info):
            label_name = label_info.get("name", "unknown")
            result["errors"].append({
                "name": name,
                "error": f"sensitivity_label_blocked:{label_name}",
                "sensitivity_blocked": True,
            })
            continue

        # Fallback: check for classification keywords in document text
        classification_keyword = check_classification_in_text(file_bytes, name)
        if classification_keyword:
            result["errors"].append({
                "name": name,
                "error": f"sensitivity_keyword_blocked:{classification_keyword}",
                "sensitivity_blocked": True,
            })
            continue

        try:
            extracted_text, page_count = extract_text(file_bytes, name)
        except (ValueError, Exception) as e:
            result["errors"].append({"name": name, "error": str(e)})
            continue

        is_tabular = ext in ("xlsx", "csv")

        if is_tabular and len(extracted_text) > SESSION_ATTR_MAX_CHARS:
            media_type = MEDIA_TYPES.get(ext, "application/octet-stream")
            result["code_interpreter_docs"].append({
                "name": name,
                "bytes": file_bytes,
                "media_type": media_type,
            })
        else:
            doc_context, context_mode = prepare_document_context(extracted_text, file_ext=ext)
            result["text_docs"].append({
                "name": name,
                "full_text": extracted_text,
                "page_count": page_count,
                "context": doc_context,
                "context_mode": context_mode,
            })

    return result
